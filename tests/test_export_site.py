"""Tests für techcore-site URL-Hilfen und Export."""

from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

import pytest
import yaml

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = PROJECT_ROOT / "scripts"
KNOWLEDGE = PROJECT_ROOT.parent / "BewerbungsAgent" / "knowledge"


def _load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


urls = _load_module("techcore_urls", SCRIPTS / "urls.py")
export_site = _load_module("techcore_export_site", SCRIPTS / "export_site.py")


@pytest.mark.parametrize(
    "folder_name, expected",
    [
        (
            "Leyton Deutschland GmbH - Inside Sales Representative (m_w_d)",
            "leyton-deutschland-gmbh-inside-sales-representative-m_w_d",
        ),
        ("Firma / Test: Sonderzeichen!!!", "firma-test-sonderzeichen"),
        ("   ", "bewerbung"),
    ],
)
def test_application_slug(folder_name: str, expected: str) -> None:
    assert urls.application_slug(folder_name) == expected


@pytest.mark.parametrize(
    "company, expected",
    [
        ("Leyton Deutschland GmbH", "leyton-deutschland-gmbh"),
        ("IFLB Laboratoriumsmedizin Berlin GmbH", "iflb-laboratoriumsmedizin-berlin-gmbh"),
        ("BMW", "bmw"),
        ("   ", "bewerbung"),
    ],
)
def test_company_slug(company: str, expected: str) -> None:
    assert urls.company_slug(company) == expected


def test_build_qr_url() -> None:
    assert urls.build_qr_url() == "https://www.techcore.de/"


def test_build_application_url() -> None:
    url = urls.build_application_url("Leyton Deutschland GmbH")
    assert url == "https://www.techcore.de/leyton-deutschland-gmbh/"


@pytest.mark.parametrize(
    "period, expected",
    [
        ({"from": "2025-05", "to": "2026-04"}, "Mai 2025 – Apr 2026"),
        ({"from": "2017-09", "to": "2020-02"}, "Sep 2017 – Feb 2020"),
        ({"from": "2021-01", "to": ""}, "Jan 2021 – heute"),
    ],
)
def test_format_period(period: dict, expected: str) -> None:
    assert export_site.format_period(period) == expected


def test_load_experiences_from_profile() -> None:
    profile_path = KNOWLEDGE / "profile.yaml"
    if not profile_path.is_file():
        pytest.skip("Profil-Wissensbasis nicht verfügbar")
    profile = export_site.load_profile(profile_path)
    experiences = export_site.load_experiences(profile, KNOWLEDGE / "experiences")
    assert len(experiences) >= 4
    assert experiences[0].period_from >= experiences[-1].period_from
    companies = {exp.company for exp in experiences}
    assert "tibas GmbH" in companies
    assert "Atzinger Verpackung GmbH" in companies


def _write_minimal_cover_letter_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("Florian Pfisterhammer")
    document.add_paragraph("Sehr geehrte Damen und Herren,")
    document.add_paragraph("mit großem Interesse bewerbe ich mich auf die ausgeschriebene Stelle.")
    document.add_paragraph("Mit freundlichen Grüßen")
    document.add_paragraph("Florian Pfisterhammer")
    document.save(path)


def _write_minimal_experience(path: Path, *, exp_id: str, company: str, role: str, period_from: str) -> None:
    path.write_text(
        yaml.safe_dump(
            {
                "id": exp_id,
                "company": company,
                "role": role,
                "location": "Teststadt",
                "period": {"from": period_from, "to": "2026-01"},
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )


def test_export_site_portfolio(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    profile_path = tmp_path / "profile.yaml"
    profile_path.write_text(
        yaml.safe_dump(
            {
                "person": {
                    "name": "Test User",
                    "title": "Test Title",
                    "summary": "Erster Absatz über Vertrieb.\n\nZweiter Absatz über Digitalisierung.",
                },
                "indexes": {"experiences": ["exp-tibas", "exp-atzinger"]},
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )

    experiences_dir = tmp_path / "experiences"
    experiences_dir.mkdir()
    _write_minimal_experience(
        experiences_dir / "exp-tibas.yaml",
        exp_id="exp-tibas",
        company="Tibas Test GmbH",
        role="Sachbearbeiter",
        period_from="2025-05",
    )
    _write_minimal_experience(
        experiences_dir / "exp-atzinger.yaml",
        exp_id="exp-atzinger",
        company="Atzinger Test GmbH",
        role="Vertrieb",
        period_from="2021-12",
    )

    sender_path = tmp_path / "sender.yaml"
    sender_path.write_text(
        yaml.safe_dump(
            {
                "sender": {
                    "zip_city": "85368 Moosburg",
                    "email": "test@example.com",
                }
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )

    apps_base = tmp_path / "apps"
    app_dir = apps_base / "Firma A - Job B"
    app_dir.mkdir(parents=True)
    _write_minimal_cover_letter_docx(app_dir / "Anschreiben.docx")
    (app_dir / "metadata.json").write_text(
        json.dumps(
            {
                "company": "Firma A",
                "job_title": "Job B",
                "folder_name": app_dir.name,
                "approved_at": "2026-07-13T00:00:00+00:00",
            }
        ),
        encoding="utf-8",
    )

    ref_base = tmp_path / "refs"
    ref_base.mkdir()
    (ref_base / "Lebenslauf 2026.docx").write_bytes(b"cv")
    (ref_base / "Atzinger Zeugnis.jpg").write_bytes(b"atzinger")
    (ref_base / "Tibas Arbeitszeugnis.jpeg").write_bytes(b"tibas")

    config_path = tmp_path / "site.config.yaml"
    config_path.write_text(
        yaml.safe_dump(
            {
                "domain": "techcore.de",
                "base_url": "https://www.techcore.de",
                "qr_landing_path": "/",
                "profile_yaml": str(profile_path),
                "experiences_dir": str(experiences_dir),
                "sender_config": str(sender_path),
                "applications_base": str(apps_base),
                "reference_base": str(ref_base),
                "require_approved": True,
                "output_dir": str(tmp_path / "dist"),
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        export_site,
        "resolve_standard_application_attachments",
        lambda: {
            "cv": ref_base / "Lebenslauf 2026.docx",
            "atzinger_certificate": ref_base / "Atzinger Zeugnis.jpg",
            "tibas_certificate": ref_base / "Tibas Arbeitszeugnis.jpeg",
        },
    )

    out = export_site.export_site(export_site.load_site_config(config_path))
    assert (out / "index.html").is_file()
    assert (out / "assets" / "site.css").is_file()
    assert (out / "assets" / "site.js").is_file()
    assert (out / "manifest.json").is_file()
    assert not (out / "bewerbung").exists()

    homepage_html = (out / "index.html").read_text(encoding="utf-8")
    assert "Test User" in homepage_html
    assert "Test Title" in homepage_html
    assert 'id="stationen"' in homepage_html
    assert "Meine Stationen" in homepage_html
    assert 'id="ueber-mich"' in homepage_html
    assert "Über mich" in homepage_html
    assert "Tibas Test GmbH" in homepage_html
    assert "Atzinger Test GmbH" in homepage_html
    assert "/assets/zeugnisse/tibas.jpeg" in homepage_html
    assert "/assets/zeugnisse/atzinger.jpg" in homepage_html
    assert "Moosburg" in homepage_html
    assert "test@example.com" in homepage_html
    assert "Lebenslauf 2026" in homepage_html
    assert "Plus+Jakarta+Sans" in homepage_html
    assert "Aktuelle Bewerbungen" not in homepage_html
    assert "Firma A" not in homepage_html
    assert "Plus Jakarta Sans" in (out / "assets" / "site.css").read_text(encoding="utf-8")

    assert (out / "assets" / "zeugnisse" / "tibas.jpeg").is_file()
    assert (out / "assets" / "zeugnisse" / "atzinger.jpg").is_file()
    assert not (out / "firma-a").exists()
    assert not (out / "assets" / "applications").exists()

    manifest = json.loads((out / "manifest.json").read_text(encoding="utf-8"))
    assert manifest["qr_url"] == "https://www.techcore.de/"
    assert manifest["structure"] == "portfolio-homepage"
    assert len(manifest["experiences"]) == 2
    assert manifest["experiences"][0]["certificate"] == "/assets/zeugnisse/tibas.jpeg"
    assert "applications" not in manifest


def test_certificate_mime_mapping() -> None:
    assert export_site._certificate_mime("pdf") == "application/pdf"
    assert export_site._certificate_mime("jpeg") == "image/jpeg"
    assert export_site._certificate_mime("jpg") == "image/jpeg"


def test_load_sender_config_location() -> None:
    sender_path = KNOWLEDGE / "docx" / "config.yaml"
    if not sender_path.is_file():
        pytest.skip("Sender-Konfiguration nicht verfügbar")
    contact = export_site.load_sender_config(sender_path)
    assert "Moosburg" in contact.location
    assert contact.email
